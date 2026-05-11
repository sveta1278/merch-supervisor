from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────── helpers ────────────────────────────────────────

def sf(run, bold=False, size=12, underline=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def sf_code(run, size=10):
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')


def pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=False, sa=0, sb=0, left=0):
    fmt = p.paragraph_format
    fmt.alignment = align
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.first_line_indent = Cm(1.25) if first else Cm(0)
    fmt.space_after = Pt(sa)
    fmt.space_before = Pt(sb)
    fmt.left_indent = Cm(left)


def para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         first=False, size=12, sa=0, sb=0, left=0, italic=False):
    p = doc.add_paragraph()
    pf(p, align=align, first=first, sa=sa, sb=sb, left=left)
    r = p.add_run(text)
    sf(r, bold=bold, size=size, italic=italic)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=False, sb=10, sa=6)
    r = p.add_run(text)
    sf(r, bold=True, size=14)


def heading2(doc, text):
    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=False, sb=8, sa=4)
    r = p.add_run(text)
    sf(r, bold=True, size=12)


def code_block(doc, sql_text):
    """Блок SQL-кода с рамкой и моноширинным шрифтом."""
    for line in sql_text.strip().split('\n'):
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.space_after = Pt(0)
        fmt.space_before = Pt(0)
        fmt.left_indent = Cm(0.5)
        # Серый фон через XML shading
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        r = p.add_run(line if line else ' ')
        sf_code(r, size=10)


def table_schema(doc, table_name, columns):
    """Таблица описания полей модели."""
    t = doc.add_table(rows=1 + len(columns), cols=3)
    t.style = 'Table Grid'

    # заголовки
    for ci, h in enumerate(['Поле', 'Тип данных', 'Описание']):
        p = t.rows[0].cells[ci].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        # голубой заголовок
        tc = t.rows[0].cells[ci]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)

    for ri, (field, dtype, desc) in enumerate(columns):
        row = t.rows[ri + 1]
        for ci, val in enumerate([field, dtype, desc]):
            p = row.cells[ci].paragraphs[0]
            p.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if ci in (0, 2) else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            r = p.add_run(val)
            r.font.name = 'Times New Roman' if ci != 0 else 'Courier New'
            r.font.size = Pt(10)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), r.font.name)

    # ширина столбцов
    for row in t.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(4.5)
        row.cells[2].width = Cm(7.5)

    doc.add_paragraph()


# ─────────────────────────── DOCUMENT ────────────────────────────────────────

doc = Document()
section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.left_margin  = Cm(3)
section.right_margin = Cm(1.5)
section.top_margin   = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ─── Заголовок раздела ────────────────────────────────────────────────────────

p = doc.add_paragraph()
pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sb=0, sa=6)
r = p.add_run('Структура базы данных веб-сервиса мониторинга мерчандайзинга')
sf(r, bold=True, size=14)

para(doc,
     'База данных системы реализована средствами Django ORM и содержит семь '
     'взаимосвязанных таблиц. В процессе разработки для локального тестирования '
     'использовалась СУБД SQLite; в производственной среде предусмотрена '
     'миграция на PostgreSQL. Названия таблиц формируются Django автоматически '
     'по схеме «имя_приложения_имя_модели» (префикс core_).',
     first=True, sa=6)

# ─── 1. Обзор таблиц ─────────────────────────────────────────────────────────

heading1(doc, '1. Перечень таблиц')

TABLES_OVERVIEW = [
    ('core_user',          'Пользователи системы (мерчандайзеры и супервайзеры)'),
    ('core_store',         'Торговые точки (магазины)'),
    ('core_sku',           'Товарные позиции (SKU)'),
    ('core_storeskuplan',  'Плановые показатели присутствия SKU в магазине'),
    ('core_visit',         'Визиты мерчандайзера в магазин'),
    ('core_visitphoto',    'Фотографии выкладки, прикреплённые к визиту'),
    ('core_visitcheckitem','Результат проверки каждой SKU в рамках визита'),
]

t_ov = doc.add_table(rows=1 + len(TABLES_OVERVIEW), cols=2)
t_ov.style = 'Table Grid'
for ci, h in enumerate(['Таблица', 'Назначение']):
    p = t_ov.rows[0].cells[ci].paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(h)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    tc = t_ov.rows[0].cells[ci]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)

for ri, (tbl, desc) in enumerate(TABLES_OVERVIEW):
    row = t_ov.rows[ri + 1]
    for ci, val in enumerate([tbl, desc]):
        p = row.cells[ci].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(val)
        r.font.name = 'Courier New' if ci == 0 else 'Times New Roman'
        r.font.size = Pt(10)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), r.font.name)
    row.cells[0].width = Cm(5.5)
    row.cells[1].width = Cm(11)

doc.add_paragraph()

# ─── 2. Описание каждой таблицы ──────────────────────────────────────────────

heading1(doc, '2. Описание полей таблиц')

heading2(doc, '2.1. core_user — Пользователи')
para(doc,
     'Расширяет стандартную модель Django AbstractUser. Добавлено поле role '
     'для разграничения прав: мерчандайзер видит только свою аналитику, '
     'суперпользователь — сводную статистику по всем сотрудникам.',
     first=True, sa=4)
table_schema(doc, 'core_user', [
    ('id',           'INTEGER PK',      'Первичный ключ (автоинкремент)'),
    ('username',     'VARCHAR(150)',     'Уникальный логин'),
    ('first_name',   'VARCHAR(150)',     'Имя'),
    ('last_name',    'VARCHAR(150)',     'Фамилия'),
    ('email',        'VARCHAR(254)',     'Электронная почта'),
    ('password',     'VARCHAR(128)',     'Хэш пароля (PBKDF2)'),
    ('role',         'VARCHAR(20)',      'Роль: merchandiser / supervisor'),
    ('is_superuser', 'BOOLEAN',         'Признак суперпользователя'),
    ('is_active',    'BOOLEAN',         'Активна ли учётная запись'),
    ('date_joined',  'DATETIME',        'Дата регистрации'),
])

heading2(doc, '2.2. core_store — Магазины')
para(doc,
     'Справочник торговых точек. Каждый магазин привязан к торговой сети '
     'и имеет уникальный адрес.',
     first=True, sa=4)
table_schema(doc, 'core_store', [
    ('id',      'INTEGER PK',  'Первичный ключ'),
    ('name',    'VARCHAR(200)','Название магазина'),
    ('network', 'VARCHAR(200)','Торговая сеть (например, Пятёрочка)'),
    ('address', 'VARCHAR(300)','Адрес торговой точки'),
])

heading2(doc, '2.3. core_sku — Товарные позиции')
para(doc,
     'Справочник SKU (Stock Keeping Unit) — товарных позиций, '
     'присутствие которых проверяется при каждом визите.',
     first=True, sa=4)
table_schema(doc, 'core_sku', [
    ('id',      'INTEGER PK',  'Первичный ключ'),
    ('name',    'VARCHAR(200)','Наименование товара'),
    ('barcode', 'VARCHAR(50)', 'Штрихкод (необязательное поле)'),
])

heading2(doc, '2.4. core_storeskuplan — Плановые показатели')
para(doc,
     'Связывает магазин и SKU: определяет, какие товары должны присутствовать '
     'на полке в данной торговой точке и в каком количестве. '
     'Комбинация (store_id, sku_id) уникальна.',
     first=True, sa=4)
table_schema(doc, 'core_storeskuplan', [
    ('id',       'INTEGER PK', 'Первичный ключ'),
    ('store_id', 'INTEGER FK', 'Ссылка на core_store.id'),
    ('sku_id',   'INTEGER FK', 'Ссылка на core_sku.id'),
    ('quantity', 'INTEGER',    'Плановое количество единиц на полке'),
])

heading2(doc, '2.5. core_visit — Визиты')
para(doc,
     'Основная транзакционная таблица. Фиксирует каждый визит мерчандайзера '
     'в магазин: дату, статус и итоговый процент выполнения плана.',
     first=True, sa=4)
table_schema(doc, 'core_visit', [
    ('id',             'INTEGER PK',  'Первичный ключ'),
    ('store_id',       'INTEGER FK',  'Ссылка на core_store.id'),
    ('user_id',        'INTEGER FK',  'Ссылка на core_user.id'),
    ('date',           'DATE',        'Дата визита'),
    ('status',         'VARCHAR(20)', 'Статус: in_progress / completed'),
    ('completion_pct', 'REAL',        'Процент выполнения плана (0–100)'),
    ('notes',          'TEXT',        'Примечания мерчандайзера'),
    ('created_at',     'DATETIME',    'Дата и время создания записи'),
])

heading2(doc, '2.6. core_visitphoto — Фотографии')
para(doc,
     'Хранит метаданные фотографий выкладки. Сами файлы сохраняются '
     'в файловой системе по пути media/photos/YYYY/MM/DD/.',
     first=True, sa=4)
table_schema(doc, 'core_visitphoto', [
    ('id',          'INTEGER PK',   'Первичный ключ'),
    ('visit_id',    'INTEGER FK',   'Ссылка на core_visit.id'),
    ('photo',       'VARCHAR(100)', 'Относительный путь к файлу'),
    ('uploaded_at', 'DATETIME',     'Дата и время загрузки'),
])

heading2(doc, '2.7. core_visitcheckitem — Чек-лист визита')
para(doc,
     'Содержит результат проверки каждой SKU в рамках визита. '
     'Записи создаются автоматически при открытии визита на основе '
     'плана магазина (core_storeskuplan).',
     first=True, sa=4)
table_schema(doc, 'core_visitcheckitem', [
    ('id',       'INTEGER PK',  'Первичный ключ'),
    ('visit_id', 'INTEGER FK',  'Ссылка на core_visit.id'),
    ('sku_id',   'INTEGER FK',  'Ссылка на core_sku.id'),
    ('status',   'VARCHAR(20)', 'Результат: present / absent / unknown'),
])

# ─── 3. SQL-запросы ──────────────────────────────────────────────────────────

heading1(doc, '3. SQL-запросы для проверки данных')

para(doc,
     'Приведённые запросы написаны для SQLite (используется в режиме '
     'разработки). При работе с PostgreSQL синтаксис идентичен, '
     'за исключением функций форматирования дат.',
     first=True, sa=6)

# 3.1
heading2(doc, '3.1. Просмотр всех таблиц в базе данных (SQLite)')
code_block(doc, """SELECT name
FROM sqlite_master
WHERE type = 'table'
ORDER BY name;""")

# 3.2
heading2(doc, '3.2. Список пользователей с ролями')
code_block(doc, """SELECT
    id,
    username,
    first_name || ' ' || last_name AS full_name,
    role,
    is_superuser,
    date_joined
FROM core_user
ORDER BY date_joined;""")

# 3.3
heading2(doc, '3.3. Список магазинов с количеством плановых SKU')
code_block(doc, """SELECT
    s.id,
    s.name        AS store_name,
    s.network,
    s.address,
    COUNT(p.id)   AS sku_count
FROM core_store s
LEFT JOIN core_storeskuplan p ON p.store_id = s.id
GROUP BY s.id, s.name, s.network, s.address
ORDER BY s.name;""")

# 3.4
heading2(doc, '3.4. Все визиты с информацией о магазине и мерчандайзере')
code_block(doc, """SELECT
    v.id,
    v.date,
    s.name                              AS store,
    u.first_name || ' ' || u.last_name  AS merchandiser,
    v.status,
    ROUND(v.completion_pct, 1)          AS completion_pct
FROM core_visit v
JOIN core_store s ON s.id = v.store_id
JOIN core_user  u ON u.id = v.user_id
ORDER BY v.date DESC, v.id DESC;""")

# 3.5
heading2(doc, '3.5. Чек-лист конкретного визита (подставить ID визита)')
code_block(doc, """SELECT
    ci.id,
    sk.name   AS sku_name,
    sk.barcode,
    ci.status
FROM core_visitcheckitem ci
JOIN core_sku sk ON sk.id = ci.sku_id
WHERE ci.visit_id = 1        -- замените 1 на нужный ID визита
ORDER BY sk.name;""")

# 3.6
heading2(doc, '3.6. Статистика визитов по магазинам')
code_block(doc, """SELECT
    s.name                             AS store,
    s.network,
    COUNT(v.id)                        AS total_visits,
    ROUND(AVG(v.completion_pct), 1)    AS avg_completion,
    ROUND(MAX(v.completion_pct), 1)    AS best_result,
    MAX(v.date)                        AS last_visit
FROM core_store s
LEFT JOIN core_visit v ON v.store_id = s.id
                       AND v.status = 'completed'
GROUP BY s.id, s.name, s.network
ORDER BY avg_completion DESC;""")

# 3.7
heading2(doc, '3.7. Статистика по каждому мерчандайзеру')
code_block(doc, """SELECT
    u.username,
    u.first_name || ' ' || u.last_name  AS full_name,
    COUNT(v.id)                         AS total_visits,
    ROUND(AVG(v.completion_pct), 1)     AS avg_completion,
    MIN(v.date)                         AS first_visit,
    MAX(v.date)                         AS last_visit
FROM core_user u
LEFT JOIN core_visit v ON v.user_id = u.id
                       AND v.status = 'completed'
WHERE u.is_superuser = 0
GROUP BY u.id, u.username, u.first_name, u.last_name
ORDER BY avg_completion DESC;""")

# 3.8
heading2(doc, '3.8. Отсутствующие товары по последнему визиту в каждый магазин')
code_block(doc, """SELECT
    s.name        AS store,
    sk.name       AS missing_sku,
    v.date        AS visit_date
FROM core_visitcheckitem ci
JOIN core_visit v  ON v.id  = ci.visit_id
JOIN core_store s  ON s.id  = v.store_id
JOIN core_sku   sk ON sk.id = ci.sku_id
WHERE ci.status = 'absent'
  AND v.date = (
      SELECT MAX(v2.date)
      FROM core_visit v2
      WHERE v2.store_id = v.store_id
        AND v2.status   = 'completed'
  )
ORDER BY s.name, sk.name;""")

# 3.9
heading2(doc, '3.9. Количество фотографий по визитам')
code_block(doc, """SELECT
    v.id,
    v.date,
    s.name            AS store,
    COUNT(p.id)       AS photo_count
FROM core_visit v
JOIN core_store s ON s.id = v.store_id
LEFT JOIN core_visitphoto p ON p.visit_id = v.id
GROUP BY v.id, v.date, s.name
ORDER BY v.date DESC;""")

# ─── 4. Рекомендуемое место вставки ──────────────────────────────────────────

heading1(doc, '4. Рекомендуемое место вставки в отчёт')
para(doc,
     'Данный материал рекомендуется включить в раздел «Проектирование '
     'информационной системы» — подраздел «Описание структуры базы данных» '
     '(после ER-диаграммы, если она присутствует в отчёте). '
     'Таблицы 2.1–2.7 образуют полное описание схемы БД; запросы из '
     'раздела 3 демонстрируют корректность реализованных связей и '
     'могут быть приведены в подразделе «Тестирование базы данных».',
     first=True, sa=4)

# ─── save ─────────────────────────────────────────────────────────────────────

out = 'БД_структура_и_запросы.docx'
doc.save(out)
print(f'Готово: {out}')
