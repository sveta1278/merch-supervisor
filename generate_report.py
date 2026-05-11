import os
import shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENSHOTS = {
    'login':        '/Users/svetlanapersutkina/Desktop/login.png',
    'registration': '/Users/svetlanapersutkina/Desktop/registration.png',
    'markets':      '/Users/svetlanapersutkina/Desktop/markets.png',
    'personal':     '/Users/svetlanapersutkina/Desktop/personal_analytics.png',
    'supervisor':   '/Users/svetlanapersutkina/Desktop/superuser_analytics.png',
}

# ─────────────────────────── helpers ────────────────────────────────────────

def set_font(run, bold=False, size=14, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True,
             space_after=0, space_before=0):
    pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25) if indent_first else Cm(0)
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)


def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent_first=True, size=14, space_after=0, space_before=0):
    p = doc.add_paragraph()
    para_fmt(p, align=align, indent_first=indent_first,
             space_after=space_after, space_before=space_before)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_before=6, space_after=6)
    run = p.add_run(text)
    set_font(run, bold=True, size=14)
    return p


def add_bullet(doc, text, size=14):
    p = doc.add_paragraph(style='List Bullet')
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=False)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_image(doc, key, caption, width=Cm(14)):
    path = SCREENSHOTS.get(key)
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
        run = p.add_run()
        run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_after=6)
    run = cap.add_run(caption)
    set_font(run, bold=False, size=12)


def add_table(doc, headers, rows, caption=''):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, bold=True, size=12)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            set_font(run, size=12)
    if caption:
        c = doc.add_paragraph()
        para_fmt(c, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
                 space_before=3, space_after=6)
        run = c.add_run(caption)
        set_font(run, size=12)
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)


def page_break(doc):
    doc.add_page_break()


# ─────────────────────────── build ──────────────────────────────────────────

def build(doc):

    # ══════════════════════════════ TITLE PAGE ═══════════════════════════════
    for text in [
        'МИНИСТЕРСТВО ЦИФРОВОГО РАЗВИТИЯ, СВЯЗИ\nИ МАССОВЫХ КОММУНИКАЦИЙ РОССИЙСКОЙ ФЕДЕРАЦИИ',
        'Ордена Трудового Красного Знамени федеральное государственное бюджетное\nобразовательное учреждение высшего образования',
    ]:
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
        run = p.add_run(text)
        set_font(run, size=12)

    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False)
    run = p.add_run('«МОСКОВСКИЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ\nСВЯЗИ И ИНФОРМАТИКИ»')
    set_font(run, bold=True, size=14)

    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_after=24)
    run = p.add_run('Кафедра «Сетевые информационные технологии и сервисы»')
    set_font(run, size=12)

    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_before=48, space_after=6)
    run = p.add_run('ОТЧЁТ')
    set_font(run, bold=True, size=14)

    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_after=6)
    run = p.add_run('по производственной практике')
    set_font(run, bold=True, size=14)

    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_after=72)
    run = p.add_run('Технологическая (проектно-технологическая)')
    set_font(run, bold=True, size=14)

    for label, val in [
        ('Выполнил:', ''),
        ('Студент группы:', '[ГРУППА]'),
        ('ФИО:', '[ФАМИЛИЯ ИМЯ ОТЧЕСТВО]'),
        ('', ''),
        ('Проверил:', ''),
        ('', 'ассистент Тришина С.В.'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if label:
            run = p.add_run(f'{label} {val}')
        else:
            run = p.add_run(val)
        set_font(run, size=14)

    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_before=48)
    run = p.add_run('Москва, 2026')
    set_font(run, size=14)

    page_break(doc)

    # ══════════════════════════════ СОДЕРЖАНИЕ ════════════════════════════════
    add_heading(doc, 'СОДЕРЖАНИЕ')
    toc_items = [
        ('Индивидуальное задание', '3'),
        ('Введение', '4'),
        ('РАЗДЕЛ 1. Анализ предметной области и постановка цели разработки', '5'),
        ('РАЗДЕЛ 2. Содержание выполненных работ во время практики', '8'),
        ('РАЗДЕЛ 3. Результаты практической деятельности и описание разработанного решения', '13'),
        ('Заключение', '17'),
        ('Список использованных источников', '18'),
        ('Приложение А. Листинги ключевых фрагментов программы', '20'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=False)
        run = p.add_run(f'{item}{"." * max(1, 70 - len(item))}{page}')
        set_font(run, size=14)
    page_break(doc)

    # ══════════════════════════ ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ ════════════════════════
    add_heading(doc, 'Индивидуальное задание')
    add_para(doc,
        'Тема практической разработки: «Разработка адаптивного веб-сервиса '
        'мониторинга мерчандайзинга с управлением визитами, фотофиксацией '
        'выкладки, чек-листами SKU и аналитикой выполнения плана».')
    add_para(doc, 'База практики: кафедра СИТиС МТУСИ.')
    add_para(doc, 'Сроки прохождения практики: с 04.04.2026 по 18.05.2026.')
    add_para(doc,
        'Цель практики заключается в закреплении теоретических знаний по '
        'разработке информационных систем и получении практических навыков '
        'проектирования, реализации, тестирования и документирования '
        'многопользовательского веб-приложения.')
    add_para(doc,
        'В ходе практики были поставлены следующие задачи: изучить предметную '
        'область контроля мерчандайзинга в торговых точках; спроектировать '
        'структуру базы данных; реализовать серверную часть на Python и Django; '
        'настроить хранение данных в PostgreSQL; реализовать механизм загрузки '
        'фотографий выкладки; организовать регистрацию и авторизацию '
        'пользователей; обеспечить разграничение доступа по ролям; разработать '
        'аналитический модуль с динамическими графиками; обеспечить адаптивность '
        'интерфейса для мобильных устройств; провести тестирование и подготовить '
        'комплект отчётной документации.')
    add_para(doc,
        'Ожидаемым результатом являлось получение работоспособного '
        'веб-приложения для супервайзеров мерчандайзинга, обеспечивающего '
        'учёт визитов в торговые точки, контроль выкладки товаров по чек-листу, '
        'фотофиксацию и аналитику выполнения плана в разрезе магазинов и '
        'сотрудников.')
    page_break(doc)

    # ══════════════════════════════ ВВЕДЕНИЕ ══════════════════════════════════
    add_heading(doc, 'Введение')
    add_para(doc,
        'Производственная практика является важным этапом формирования '
        'профессиональных компетенций будущего специалиста по направлению '
        '«Информационные системы и технологии». В рамках практики студент '
        'получает возможность применить теоретические знания в процессе создания '
        'реального программного продукта, включающего как серверную логику, так '
        'и механизмы взаимодействия с пользователем и базой данных.')
    add_para(doc,
        'В качестве объекта практической разработки была выбрана '
        'автоматизированная система контроля мерчандайзинга. Подобные системы '
        'широко применяются в организациях розничной торговли для планирования '
        'и контроля выкладки товаров, регистрации визитов полевых сотрудников '
        'и сбора аналитики по торговым точкам. Актуальность темы обусловлена '
        'необходимостью цифровизации полевых процессов и снижения доли ручного '
        'труда при контроле ассортимента.')
    add_para(doc,
        'В ходе практики был разработан адаптивный веб-сервис на основе языка '
        'Python, фреймворка Django 5.x и СУБД PostgreSQL. В системе реализованы '
        'регистрация и авторизация пользователей, разграничение доступа по ролям, '
        'управление торговыми точками и ассортиментными позициями (SKU), '
        'создание и ведение визитов, фотофиксация выкладки, чек-листы проверки '
        'наличия товара, а также двухуровневая аналитика: личная статистика '
        'мерчандайзера и сводная аналитика для администратора.')
    add_para(doc,
        'Целью отчёта является описание выполненной работы, архитектуры '
        'приложения, используемых технологий, этапов реализации и достигнутых '
        'результатов. В документе также приведены схема базы данных, '
        'пользовательский поток, скриншоты интерфейса и листинги ключевых '
        'фрагментов программного кода.')
    page_break(doc)

    # ══════════════════════════════ РАЗДЕЛ 1 ══════════════════════════════════
    add_heading(doc, 'РАЗДЕЛ 1. Анализ предметной области и постановка цели разработки')
    add_para(doc,
        'Объектом исследования является процесс контроля выкладки товаров в '
        'торговых точках розничной сети. В типовой ситуации мерчандайзер '
        'посещает закреплённые магазины, проверяет наличие плановых SKU '
        'на полках, фотографирует выкладку и фиксирует результаты. '
        'Супервайзер анализирует итоги по всем сотрудникам и магазинам.')
    add_para(doc,
        'Предметом исследования выступают методы и программные средства '
        'автоматизации учёта визитов, управления чек-листами и формирования '
        'аналитической отчётности в многопользовательской информационной системе.')
    add_para(doc,
        'Цель разработки состоит в создании программного приложения, '
        'позволяющего централизованно регистрировать визиты, контролировать '
        'план размещения SKU, хранить фотодоказательства выкладки, а также '
        'обеспечивать разграничение доступа по ролям и динамическую аналитику.')
    add_para(doc,
        'Для решения поставленной задачи была выбрана архитектура '
        'веб-приложения на основе Django MVT (Model-View-Template). Такой подход '
        'позволяет объединить средства маршрутизации HTTP-запросов, работу с '
        'шаблонами интерфейса, объектно-реляционное отображение через Django ORM, '
        'механизмы аутентификации и удобные средства администрирования данных.')
    add_para(doc,
        'Ключевые требования к системе сформулированы следующим образом:')
    for req in [
        'регистрация и аутентификация пользователей;',
        'разделение прав доступа по ролям: мерчандайзер и администратор (superuser);',
        'управление справочниками торговых точек (Store) и товарных позиций (SKU);',
        'формирование плана SKU для каждого магазина (StoreSKUPlan);',
        'создание визитов с автоматическим генерированием чек-листа (VisitCheckItem);',
        'загрузка и хранение фотографий выкладки (VisitPhoto);',
        'отметка статуса SKU (есть / нет / не проверено) с поддержкой AJAX;',
        'расчёт процента выполнения плана при завершении визита;',
        'личная аналитика мерчандайзера и сводная аналитика администратора;',
        'адаптивный интерфейс Bootstrap 5 с нижней навигацией для мобильных устройств;',
        'хранение данных в PostgreSQL, поддержка SQLite для разработки.',
    ]:
        add_bullet(doc, req)

    add_para(doc,
        'На основании изучения предметной области было установлено, что '
        'приложение включает контроллеры авторизации и управления визитами, '
        'доменные модели сущностей, формы ввода данных, административный '
        'интерфейс и HTML-шаблоны. Сводная характеристика основных '
        'компонентов разработанной системы представлена в таблице 1.',
        space_before=6)

    add_table(doc,
        ['Компонент', 'Назначение', 'Технологии'],
        [
            ['Модели данных',    'Описание сущностей и хранение в БД',             'Django ORM, PostgreSQL'],
            ['Представления',   'Обработка HTTP-запросов, бизнес-логика',          'Django Views'],
            ['Шаблоны',         'HTML-страницы с адаптивным интерфейсом',          'Django Templates, Bootstrap 5'],
            ['Формы',           'Валидация и обработка пользовательского ввода',   'Django Forms'],
            ['Аутентификация',  'Вход, выход, регистрация, ограничение доступа',   'Django Auth'],
            ['Загрузка файлов', 'Сохранение и отображение фото выкладки',          'Pillow, Django FileField'],
            ['Аналитика',       'Графики и KPI по визитам и магазинам',            'Chart.js 4, Django ORM'],
            ['Администрирование','Управление справочниками через браузер',         'Django Admin'],
        ],
        caption='Таблица 1 – Основные компоненты разработанной системы')

    add_para(doc, 'Интерфейс входа пользователя в систему представлен на рисунке 1.',
             space_before=6)
    add_image(doc, 'login',
              'Рисунок 1 – Страница входа в систему', width=Cm(10))
    page_break(doc)

    # ══════════════════════════════ РАЗДЕЛ 2 ══════════════════════════════════
    add_heading(doc, 'РАЗДЕЛ 2. Содержание выполненных работ во время практики')
    add_para(doc,
        'В качестве основного языка разработки серверной части приложения '
        'использовался Python версии 3.12. Разработка, отладка и запуск '
        'приложения выполнялись в среде Visual Studio Code с расширениями '
        'для Python и Django. Применение данной среды позволило упростить '
        'навигацию по структуре проекта, работу с виртуальным окружением, '
        'запуск тестового сервера и анализ исходного кода.')
    add_para(doc,
        'Практическая работа была организована поэтапно. На первом этапе '
        'выполнена постановка задачи: изучена предметная область мерчандайзинга, '
        'определён перечень сущностей системы, составлена схема базы данных '
        'и сформированы требования к интерфейсу.')
    add_para(doc,
        'На втором этапе проведено проектирование структуры данных. '
        'Для хранения данных использована СУБД PostgreSQL, обеспечивающая '
        'надёжное хранение информации о пользователях, магазинах, SKU, '
        'визитах и фотографиях. В качестве ORM применяется встроенный '
        'механизм Django.')

    add_para(doc, 'Схема базы данных системы включает следующие сущности:',
             space_before=6)
    for entity in [
        'User (AbstractUser) — пользователь системы с полем role = «supervisor»;',
        'Store — торговая точка: название, адрес, торговая сеть, признак активности;',
        'SKU — товарная позиция: название, уникальный штрихкод, категория;',
        'StoreSKUPlan — план выкладки: связь магазина и SKU (unique_together);',
        'Visit — визит: супервайзер, магазин, дата, статус, процент выполнения, заметки;',
        'VisitPhoto — фото выкладки: файл изображения, комментарий, дата загрузки;',
        'VisitCheckItem — строка чек-листа: визит, SKU, статус (present / absent / not_checked).',
    ]:
        add_bullet(doc, entity)

    add_para(doc,
        'Схема связей между сущностями базы данных приведена в таблице 2.',
        space_before=6)
    add_table(doc,
        ['Сущность', 'Поля', 'Связи'],
        [
            ['User',            'id, username, password, role',                    '—'],
            ['Store',           'id, name, address, network, is_active',           '—'],
            ['SKU',             'id, name, barcode, category',                     '—'],
            ['StoreSKUPlan',    'id, store_id, sku_id',                            'FK→Store, FK→SKU'],
            ['Visit',           'id, supervisor_id, store_id, date, status,\ncompletion_pct, notes',
             'FK→User, FK→Store'],
            ['VisitPhoto',      'id, visit_id, image, comment, created_at',        'FK→Visit'],
            ['VisitCheckItem',  'id, visit_id, sku_id, status',                    'FK→Visit, FK→SKU'],
        ],
        caption='Таблица 2 – Схема базы данных')

    add_para(doc,
        'Пользовательский поток (User Flow) в системе реализован следующим образом. '
        'Новый пользователь переходит на страницу регистрации, вводит имя, фамилию, '
        'логин и пароль и получает учётную запись с ролью мерчандайзера. После входа '
        'он попадает на дашборд с визитами за текущий день. Для начала работы '
        'пользователь переходит в раздел «Магазины», выбирает нужную торговую точку '
        'и нажимает «Визит». Система автоматически создаёт запись визита и '
        'генерирует чек-лист со всеми SKU из плана данного магазина.',
        space_before=6)
    add_para(doc,
        'Далее на странице визита мерчандайзер работает с тремя вкладками: '
        '«Фото» — загружает снимки выкладки с комментариями; «Чек-лист» — '
        'отмечает статус каждого SKU (✓ Есть / ✗ Нет / ? Не проверено); '
        '«Детали» — вводит заметки и завершает визит. При завершении '
        'рассчитывается процент выполнения плана и открывается страница отчёта '
        'с итоговой статистикой и галереей фотографий.')
    add_para(doc,
        'Форма регистрации нового пользователя приведена на рисунке 2.')
    add_image(doc, 'registration',
              'Рисунок 2 – Форма регистрации пользователя', width=Cm(10))

    add_para(doc, 'Страница со списком торговых точек показана на рисунке 3.',
             space_before=6)
    add_image(doc, 'markets',
              'Рисунок 3 – Список торговых точек с поиском и кнопками создания визита',
              width=Cm(14))

    add_para(doc,
        'Отдельное внимание было уделено реализации модуля аналитики. '
        'Система обеспечивает двухуровневый доступ: обычный мерчандайзер '
        'видит только свою личную статистику — количество визитов, средний '
        'процент выполнения, разбивку по магазинам и динамику во времени. '
        'Администратор (superuser) получает доступ к сводной аналитике по '
        'всем сотрудникам, включая рейтинг мерчандайзеров и сравнительные '
        'графики по торговым точкам.',
        space_before=6)

    add_para(doc,
        'За период практики были выполнены следующие виды работ:')
    for work in [
        'анализ предметной области и формирование требований к системе;',
        'проектирование схемы базы данных (7 связанных моделей);',
        'настройка окружения: Python 3.12, Django 5.2, виртуальное окружение, .env;',
        'разработка моделей данных и миграций Django;',
        'реализация представлений (views) и URL-маршрутизации;',
        'создание форм с валидацией, в том числе формы регистрации и загрузки фото;',
        'разработка административного интерфейса с InlineAdmin;',
        'реализация адаптивных HTML-шаблонов на Bootstrap 5 с нижней навигацией;',
        'интеграция Chart.js 4 для построения интерактивных графиков аналитики;',
        'реализация AJAX-обновления статусов SKU в чек-листе через Fetch API;',
        'тестирование пользовательских сценариев и исправление обнаруженных недочётов;',
        'подготовка отчётной документации по результатам практики.',
    ]:
        add_bullet(doc, work)

    add_para(doc, 'Основные этапы выполнения работ сведены в таблицу 3.',
             space_before=6)
    add_table(doc,
        ['Этап', 'Содержание работ', 'Результат', 'Срок'],
        [
            ['1', 'Изучение предметной\nобласти и постановка задачи',
             'Требования к системе,\nперечень сущностей', '1 неделя'],
            ['2', 'Проектирование БД и\nструктуры проекта',
             'Схема данных, настроенное\nокружение', '1–2 недели'],
            ['3', 'Реализация моделей,\nпредставлений, форм',
             'Рабочий CRUD для визитов,\nчек-листов, фото', '2–4 недели'],
            ['4', 'Разработка аналитики\nи интерфейса',
             'Графики Chart.js,\nадаптивный дизайн', '4–5 недели'],
            ['5', 'Тестирование и\nдокументирование',
             'Отчёт, скриншоты,\nфинальная версия кода', '6 неделя'],
        ],
        caption='Таблица 3 – Этапы выполнения работ в период практики')
    page_break(doc)

    # ══════════════════════════════ РАЗДЕЛ 3 ══════════════════════════════════
    add_heading(doc, 'РАЗДЕЛ 3. Результаты практической деятельности и описание разработанного решения')
    add_para(doc,
        'Результатом практики стал работоспособный веб-сервис мониторинга '
        'мерчандайзинга, предназначенный для полевых сотрудников торговых сетей. '
        'Приложение построено по архитектурному шаблону MVT (Model-View-Template) '
        'и включает серверную часть на Django, реляционную базу данных, '
        'адаптивный пользовательский интерфейс и модуль аналитики.')
    add_para(doc,
        'На уровне предметной модели реализованы семь взаимосвязанных сущностей, '
        'описывающих полный жизненный цикл визита: от планирования ассортимента '
        'до формирования итогового отчёта. Модель User расширяет стандартный '
        'AbstractUser полем роли, что обеспечивает основу для разграничения прав. '
        'Модели Store и SKU формируют справочники торговых точек и товарных '
        'позиций. StoreSKUPlan связывает магазины с плановым ассортиментом. '
        'Visit фиксирует каждый выезд сотрудника, VisitPhoto хранит '
        'фотодоказательства, VisitCheckItem — построчный результат проверки.')
    add_para(doc,
        'С точки зрения безопасности система обеспечивает обязательную '
        'аутентификацию пользователя и разграничение доступа. Все представления '
        'защищены декоратором @login_required. Мерчандайзер видит исключительно '
        'свои визиты — фильтрация по полю supervisor=request.user применяется '
        'на уровне запросов к базе данных. CSRF-токен включён во все формы. '
        'При загрузке файлов проверяется MIME-тип: принимаются только изображения. '
        'Администратор (superuser) получает расширенный доступ к сводной '
        'аналитике по всем сотрудникам.')
    add_para(doc,
        'Личная аналитика мерчандайзера представлена на рисунке 4.')
    add_image(doc, 'personal',
              'Рисунок 4 – Личная аналитика мерчандайзера с KPI и графиками',
              width=Cm(15))

    add_para(doc,
        'Раздел аналитики для администратора включает шесть интерактивных '
        'графиков, построенных с использованием библиотеки Chart.js 4: '
        'столбчатая диаграмма визитов по дням, линейный график динамики '
        'выполнения плана, горизонтальные гистограммы рейтинга мерчандайзеров '
        'по количеству визитов и по проценту выполнения, горизонтальная '
        'гистограмма по торговым точкам, а также круговая диаграмма '
        'распределения визитов по диапазонам выполнения (0–25%, 25–50%, '
        '50–75%, 75–100%). Цветовая индикация на графиках отражает уровень '
        'выполнения: зелёный — более 80%, жёлтый — 50–80%, красный — менее 50%.',
        space_before=6)
    add_para(doc,
        'Сводная аналитика администратора представлена на рисунке 5.')
    add_image(doc, 'supervisor',
              'Рисунок 5 – Аналитика администратора по всем мерчандайзерам',
              width=Cm(15))

    add_para(doc,
        'В ходе тестирования были проверены основные сценарии работы системы:',
        space_before=6)
    for scenario in [
        'регистрация нового пользователя и вход в систему;',
        'поиск магазина по названию и создание нового визита;',
        'автоматическое формирование чек-листа из плана магазина;',
        'загрузка фотографий выкладки с комментариями;',
        'отметка статусов SKU через кнопки (без перезагрузки страницы через AJAX);',
        'сохранение заметок и завершение визита с расчётом процента выполнения;',
        'просмотр итогового отчёта визита с прогресс-баром и галереей фото;',
        'фильтрация аналитики по периоду (date_from, date_to);',
        'проверка изоляции данных: мерчандайзер видит только свои визиты;',
        'проверка корректной работы на экране 375px (мобильный вид).',
    ]:
        add_bullet(doc, scenario)

    add_para(doc,
        'По результатам проверки было установлено, что приложение корректно '
        'запускается, сохраняет данные в базе, поддерживает основные ролевые '
        'сценарии и корректно отображается на мобильных устройствах. '
        'В перспективе систему можно расширить push-уведомлениями при создании '
        'визита, экспортом отчётов в Excel, модулем планирования маршрутов '
        'и интеграцией с геолокацией для подтверждения присутствия в магазине.',
        space_before=6)
    page_break(doc)

    # ══════════════════════════════ ЗАКЛЮЧЕНИЕ ════════════════════════════════
    add_heading(doc, 'Заключение')
    add_para(doc,
        'В период прохождения производственной практики был разработан '
        'адаптивный веб-сервис мониторинга мерчандайзинга на основе Python 3.12, '
        'фреймворка Django 5.2 и СУБД PostgreSQL. Поставленная цель по созданию '
        'работоспособного программного приложения с полным набором заявленных '
        'функций достигнута.')
    add_para(doc,
        'В ходе работы были закреплены знания по объектно-ориентированному '
        'программированию на Python, веб-разработке с использованием Django MVT, '
        'проектированию реляционных баз данных и построению ORM-моделей, '
        'разработке адаптивных интерфейсов на Bootstrap 5, интеграции '
        'JavaScript-библиотек визуализации данных, а также организации '
        'ролевой модели доступа.')
    add_para(doc,
        'Разработанная система обеспечивает регистрацию и авторизацию '
        'пользователей, планирование и проведение визитов в торговые точки, '
        'фотофиксацию выкладки, построчную проверку наличия SKU по чек-листу, '
        'автоматический расчёт процента выполнения плана и централизованное '
        'хранение данных. Модуль аналитики реализует двухуровневую систему '
        'отчётности с интерактивными графиками на Chart.js.')
    add_para(doc,
        'Полученные результаты подтверждают практическую значимость '
        'выполненной работы. Созданное приложение может использоваться как '
        'демонстрационный прототип системы контроля мерчандайзинга и являться '
        'основой для дальнейшего развития в рамках выпускной '
        'квалификационной работы.')
    page_break(doc)

    # ══════════════════════════ СПИСОК ИСТОЧНИКОВ ════════════════════════════
    add_heading(doc, 'Список использованных источников')
    sources = [
        'Python 3.12 Documentation [Электронный ресурс]. — URL: https://docs.python.org/3.12/ (дата обращения: 10.05.2026).',
        'Django 5.x Documentation [Электронный ресурс]. — URL: https://docs.djangoproject.com/en/5.2/ (дата обращения: 10.05.2026).',
        'Django Authentication Documentation [Электронный ресурс]. — URL: https://docs.djangoproject.com/en/5.2/topics/auth/ (дата обращения: 10.05.2026).',
        'PostgreSQL Documentation [Электронный ресурс]. — URL: https://www.postgresql.org/docs/current/index.html (дата обращения: 10.05.2026).',
        'Bootstrap 5 Documentation [Электронный ресурс]. — URL: https://getbootstrap.com/docs/5.3/ (дата обращения: 10.05.2026).',
        'Chart.js Documentation [Электронный ресурс]. — URL: https://www.chartjs.org/docs/latest/ (дата обращения: 10.05.2026).',
        'Pillow (Python Imaging Library Fork) Documentation [Электронный ресурс]. — URL: https://pillow.readthedocs.io/en/stable/ (дата обращения: 10.05.2026).',
        'python-dotenv Documentation [Электронный ресурс]. — URL: https://saurabh-kumar.com/python-dotenv/ (дата обращения: 10.05.2026).',
        'Django ORM QuerySet API Reference [Электронный ресурс]. — URL: https://docs.djangoproject.com/en/5.2/ref/models/querysets/ (дата обращения: 10.05.2026).',
        'MDN Web Docs — Fetch API [Электронный ресурс]. — URL: https://developer.mozilla.org/en-US/docs/Web/API/Fetch_API (дата обращения: 10.05.2026).',
        'Лутц М. Изучаем Python. — 5-е изд. — Москва: O\'Reilly Media / ДМК Пресс, 2020.',
        'Рамальо Л. Python. Лучшие практики и инструменты. — 2-е изд. — Москва: Питер, 2023.',
        'Форман Д. Django для профессионалов. Веб-разработка на Python. — Москва: Вильямс, 2022.',
        'Грейф Т. Bootstrap. Адаптивная веб-разработка. — Москва: ДМК Пресс, 2021.',
        'Мартин Р. Чистый код: создание, анализ и рефакторинг. — Санкт-Петербург: Питер, 2019.',
    ]
    for i, src in enumerate(sources, 1):
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=False)
        p.paragraph_format.left_indent = Cm(1.25)
        run = p.add_run(f'{i}. {src}')
        set_font(run, size=14)
    page_break(doc)

    # ══════════════════════════ ПРИЛОЖЕНИЕ А ══════════════════════════════════
    add_heading(doc, 'Приложение А. Листинги ключевых фрагментов программы')

    add_para(doc, 'Листинг 1 – Определение модели Visit',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_before=6)
    add_code(doc, """class Visit(models.Model):
    STATUS_CHOICES = [
        ('in_progress', 'В процессе'),
        ('completed',   'Завершён'),
    ]
    supervisor    = models.ForeignKey(User, on_delete=models.CASCADE,
                                      related_name='visits')
    store         = models.ForeignKey(Store, on_delete=models.CASCADE)
    date          = models.DateField(auto_now_add=True)
    status        = models.CharField(max_length=20, choices=STATUS_CHOICES,
                                     default='in_progress')
    completion_pct = models.FloatField(null=True, blank=True)
    notes         = models.TextField(blank=True)""")

    add_para(doc, 'Листинг 2 – Создание визита и автоматическое формирование чек-листа',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_before=6)
    add_code(doc, """@login_required
def visit_new(request):
    store = get_object_or_404(Store, pk=request.GET.get('store'),
                              is_active=True)
    if request.method == 'POST':
        visit = Visit.objects.create(supervisor=request.user, store=store)
        plan_skus = StoreSKUPlan.objects.filter(store=store).select_related('sku')
        VisitCheckItem.objects.bulk_create([
            VisitCheckItem(visit=visit, sku=p.sku, status='not_checked')
            for p in plan_skus
        ])
        return redirect('visit_detail', pk=visit.pk)
    return render(request, 'visits/new.html', {'store': store})""")

    add_para(doc, 'Листинг 3 – Завершение визита с расчётом процента выполнения',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_before=6)
    add_code(doc, """elif action == 'complete_visit':
    items = visit.check_items.all()
    total   = items.count()
    present = items.filter(status='present').count()
    visit.completion_pct = (present / total * 100) if total > 0 else 0
    visit.status = 'completed'
    visit.save()
    return redirect('visit_report', pk=pk)""")

    add_para(doc, 'Листинг 4 – Агрегация статистики по мерчандайзерам для администратора',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_before=6)
    add_code(doc, """supervisor_stats = (
    visits_qs
    .values('supervisor__id', 'supervisor__first_name',
            'supervisor__last_name', 'supervisor__username')
    .annotate(
        visit_count = Count('id'),
        avg_pct     = Avg('completion_pct'),
        last_visit  = Max('date'),
    )
    .order_by('-visit_count')
)""")

    add_para(doc, 'Листинг 5 – AJAX-обновление статуса SKU в чек-листе (JavaScript)',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, indent_first=False,
             space_before=6)
    add_code(doc, """document.querySelectorAll('.check-form').forEach(function(form) {
  form.addEventListener('submit', function(e) {
    e.preventDefault();
    var data = new FormData(form);
    fetch('', {
      method: 'POST',
      body: data,
      headers: {'X-CSRFToken': data.get('csrfmiddlewaretoken')},
    }).then(function(resp) {
      if (resp.ok) {
        var itemId = data.get('item_id');
        var status = data.get('status');
        updateButtonStates(itemId, status);
        updateCounter(status, itemId);
      }
    });
  });
});""")


# ─────────────────────────── main ────────────────────────────────────────────

doc = Document()
build(doc)
out = 'Отчет_по_практике_Мерчандайзинг.docx'
doc.save(out)
print(f'Готово: {out}')
