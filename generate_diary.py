from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────── helpers ────────────────────────────────────────

def sf(run, bold=False, size=14, underline=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=True,
       sa=0, sb=0, left=0):
    fmt = p.paragraph_format
    fmt.alignment = align
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.first_line_indent = Cm(1.25) if first else Cm(0)
    fmt.space_after  = Pt(sa)
    fmt.space_before = Pt(sb)
    fmt.left_indent  = Cm(left)


def para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         first=True, size=14, sa=0, sb=0, left=0,
         underline=False, italic=False):
    p = doc.add_paragraph()
    pf(p, align=align, first=first, sa=sa, sb=sb, left=left)
    r = p.add_run(text)
    sf(r, bold=bold, size=size, underline=underline, italic=italic)
    return p


def heading(doc, text, size=14):
    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sb=6, sa=6)
    r = p.add_run(text)
    sf(r, bold=True, size=size)


def signature_line(doc, label, hint='(подпись)', date=''):
    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=False, sa=0)
    r = p.add_run(label)
    sf(r, bold=True, size=14)

    p2 = doc.add_paragraph()
    pf(p2, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=2)
    r2 = p2.add_run('_' * 40)
    sf(r2, size=14)

    p3 = doc.add_paragraph()
    pf(p3, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=6)
    r3 = p3.add_run(hint)
    sf(r3, italic=True, size=12)

    if date:
        p4 = doc.add_paragraph()
        pf(p4, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=6)
        r4 = p4.add_run(date)
        sf(r4, size=14)


def pb(doc):
    doc.add_page_break()


def add_row_to_table(table, row_idx, cells_data, sizes=None, bold_col=None):
    row = table.rows[row_idx].cells
    for ci, val in enumerate(cells_data):
        p = row[ci].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r = p.add_run(val)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(sizes[ci] if sizes else 12)
        r.bold = (bold_col is not None and ci == bold_col)


# ─────────────────────────── PAGES ──────────────────────────────────────────

def title_page(doc):
    para(doc, 'МИНИСТЕРСТВО ЦИФРОВОГО РАЗВИТИЯ, СВЯЗИ И МАССОВЫХ\nКОММУНИКАЦИЙ РОССИЙСКОЙ ФЕДЕРАЦИИ',
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, size=12)
    para(doc, 'Ордена Трудового Красного Знамени федеральное государственное бюджетное\nобразовательное учреждение высшего образования',
         align=WD_ALIGN_PARAGRAPH.CENTER, first=False, size=12, underline=True)
    para(doc, '«Московский технический университет связи и информатики»\n(МТУСИ)',
         align=WD_ALIGN_PARAGRAPH.CENTER, first=False, size=12, underline=True, sb=24)

    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sb=48, sa=6)
    r = p.add_run('Д Н Е В Н И К')
    sf(r, bold=True, size=14)

    para(doc, 'по производственной практике\nТехнологическая (проектно-технологическая)',
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, underline=True, sa=4)
    para(doc, '(вид и тип практики)',
         align=WD_ALIGN_PARAGRAPH.CENTER, first=False, italic=True, size=12, sa=48)

    para(doc, 'студентки [ФАМИЛИЯ ИМЯ ОТЧЕСТВО]',
         align=WD_ALIGN_PARAGRAPH.LEFT, first=False, sa=6)
    para(doc, 'группы [ГРУППА]',
         align=WD_ALIGN_PARAGRAPH.LEFT, first=False, sa=72)

    para(doc, 'Москва, 2026',
         align=WD_ALIGN_PARAGRAPH.CENTER, first=False)
    pb(doc)


def direction_page(doc):
    para(doc,
         'Ордена Трудового Красного Знамени федеральное государственное бюджетное '
         'образовательное учреждение высшего образования «Московский технический '
         'университет связи и информатики» (МТУСИ) на основании договора'
         + '_' * 40,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, first=False, sa=6)
    para(doc, '_' * 80, first=False, sa=12)

    para(doc, 'Направляет [ФАМИЛИЯ ИМЯ ОТЧЕСТВО]' + '_' * 20,
         first=False, sa=0)
    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=6)
    r = p.add_run('(фамилия, имя, отчество)')
    sf(r, italic=True, size=12)

    para(doc, 'Группы [ГРУППА]' + '_' * 50, first=False, sa=6)

    p2 = doc.add_paragraph()
    pf(p2, first=False, sa=2)
    r2 = p2.add_run('для прохождения практики ')
    sf(r2, size=14)
    r3 = p2.add_run('производственной практики Технологическая (проектно-технологическая)')
    sf(r3, bold=True, underline=True, size=14)

    p3 = doc.add_paragraph()
    pf(p3, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=4)
    r4 = p3.add_run('(вид и тип практики)')
    sf(r4, italic=True, size=12)

    para(doc, 'на кафедре СИТиС МТУСИ' + '_' * 40, first=False, sa=0)
    p4 = doc.add_paragraph()
    pf(p4, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=12)
    r5 = p4.add_run('(наименование базы практики)')
    sf(r5, italic=True, size=12)

    para(doc, 'Срок   с 04.04.2026 г. по 18.05.2026 г.', first=False, sa=6)
    para(doc, 'Период практики: 04.04.2026–18.05.2026.', first=False, sa=18)

    p5 = doc.add_paragraph()
    pf(p5, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=0)
    r6 = p5.add_run('Декан ЦЗОПБ')
    sf(r6, size=14)
    p6 = doc.add_paragraph()
    pf(p6, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=2)
    r7 = p6.add_run('_' * 20)
    sf(r7, size=14)
    p7 = doc.add_paragraph()
    pf(p7, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=18)
    r8 = p7.add_run('(подпись)\nПечать МТУСИ')
    sf(r8, italic=True, size=12)

    para(doc, 'Прибыл в организацию "04" апреля 2026 г.', first=False, sa=6)
    p8 = doc.add_paragraph()
    pf(p8, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=2)
    r9 = p8.add_run('_' * 30)
    sf(r9, size=14)
    p9 = doc.add_paragraph()
    pf(p9, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=12)
    r10 = p9.add_run('(подпись)')
    sf(r10, italic=True, size=12)

    para(doc, 'Выбыл из организации "18" мая 2026 г.', first=False, sa=6)
    p10 = doc.add_paragraph()
    pf(p10, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=2)
    r11 = p10.add_run('_' * 30)
    sf(r11, size=14)
    p11 = doc.add_paragraph()
    pf(p11, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=12)
    r12 = p11.add_run('(подпись)')
    sf(r12, italic=True, size=12)

    p12 = doc.add_paragraph()
    pf(p12, first=False, sa=0)
    r13 = p12.add_run('Печать (организации)')
    sf(r13, bold=True, size=14)
    pb(doc)


def diary_table(doc):
    heading(doc, 'Д Н Е В Н И К')

    ENTRIES = [
        ('04.04.2026',
         'Кафедра СИТиС МТУСИ. Прошла инструктаж по технике безопасности, '
         'ознакомилась с целями практики, шаблонами отчётности и темой '
         'индивидуального задания.'),
        ('06.04.2026',
         'Рабочее место разработчика. Изучила предметную область контроля '
         'мерчандайзинга: процесс визитов в торговые точки, проверка '
         'ассортимента, фотофиксация выкладки.'),
        ('08.04.2026',
         'Рабочее место разработчика. Провела анализ требований, определила '
         'перечень сущностей системы: пользователь, магазин, SKU, план, '
         'визит, фото, чек-лист.'),
        ('10.04.2026',
         'Рабочее место разработчика. Установила среду разработки: Python 3.12, '
         'Django 5.2, виртуальное окружение, зависимости из requirements.txt. '
         'Создала структуру проекта.'),
        ('13.04.2026',
         'Рабочее место разработчика. Разработала модели данных User, Store, '
         'SKU и StoreSKUPlan, настроила параметры AUTH_USER_MODEL и '
         'выполнила первые миграции.'),
        ('15.04.2026',
         'Рабочее место разработчика. Реализовала модели Visit, VisitPhoto и '
         'VisitCheckItem. Проверила связи внешних ключей и ограничение '
         'unique_together для StoreSKUPlan.'),
        ('17.04.2026',
         'Рабочее место разработчика. Настроила Django Admin: зарегистрировала '
         'все модели, реализовала StoreSKUPlanInline и VisitCheckItemInline, '
         'добавила поиск и фильтры.'),
        ('20.04.2026',
         'Рабочее место разработчика. Разработала представление дашборда '
         '(список визитов за текущий день) и страницу списка магазинов '
         'с поиском по GET-параметру.'),
        ('22.04.2026',
         'Рабочее место разработчика. Реализовала создание визита с '
         'автоматическим формированием чек-листа через bulk_create '
         'из плана SKU магазина.'),
        ('24.04.2026',
         'Рабочее место разработчика. Разработала страницу визита с тремя '
         'Bootstrap-вкладками: «Фото», «Чек-лист», «Детали». '
         'Настроила переключение вкладок через GET-параметр.'),
        ('27.04.2026',
         'Рабочее место разработчика. Реализовала загрузку фотографий '
         'выкладки: форма VisitPhotoForm с валидацией MIME-типа, '
         'автосабмит по выбору файла, отображение галереи.'),
        ('29.04.2026',
         'Рабочее место разработчика. Разработала чек-лист SKU с тремя '
         'кнопками-тоглами (✓/✗/?) и AJAX-обновлением статуса через '
         'Fetch API без перезагрузки страницы.'),
        ('01.05.2026',
         'Рабочее место разработчика. Реализовала завершение визита: расчёт '
         'completion_pct = present/total × 100, обновление статуса, '
         'редирект на страницу итогового отчёта.'),
        ('04.05.2026',
         'Рабочее место разработчика. Разработала модуль аналитики: '
         'агрегация по магазинам через Django ORM (Count, Avg, Max), '
         'интеграция Chart.js 4 для четырёх типов графиков.'),
        ('06.05.2026',
         'Рабочее место разработчика. Реализовала страницу регистрации '
         'мерчандайзера (UserRegistrationForm), обновила страницу входа '
         'со ссылкой «Зарегистрироваться».'),
        ('08.05.2026',
         'Рабочее место разработчика. Реализовала двухуровневую аналитику: '
         'личная статистика для мерчандайзера и сводная по всем '
         'сотрудникам для суперпользователя.'),
        ('11.05.2026',
         'Рабочее место разработчика. Исправила обработку выхода из системы '
         '(LogoutView в Django 5 требует POST), добавила форму с '
         'CSRF-токеном в navbar.'),
        ('13.05.2026',
         'Рабочее место разработчика. Провела тестирование пользовательских '
         'сценариев: регистрация, создание визита, отметка SKU, '
         'загрузка фото, завершение, аналитика.'),
        ('14.05.2026',
         'Рабочее место разработчика. Заполнила дневник практики, '
         'индивидуальное задание, рабочий план и отзыв. '
         'Оформила отчёт по разработанному проекту.'),
        ('15.05.2026',
         'Рабочее место разработчика. Завершила подготовку отчёта по практике, '
         'дневника, индивидуального задания и сопроводительных материалов, '
         'выполнила итоговую проверку документов.'),
        ('18.05.2026',
         'Рабочее место разработчика. Завершила прохождение практики, '
         'согласовала и передала комплект отчётной документации, '
         'подготовила документы к итоговому подписанию.'),
    ]

    table = doc.add_table(rows=1 + len(ENTRIES), cols=3)
    table.style = 'Table Grid'

    # Header
    hdrs = ['Дата выполнения\nработы',
            'Рабочее место и краткое содержание выполняемых работ\n'
            '(вписываются конкретные работы, выполняемые\n'
            'обучающимся на рабочем месте)',
            'Подпись руководителя\nпрактики']
    for ci, h in enumerate(hdrs):
        cell = table.rows[0].cells[ci]
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(3)

    for ri, (date, content) in enumerate(ENTRIES):
        row = table.rows[ri + 1]
        for ci, val in enumerate([date, content, '__________']):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 2)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            r = p.add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    pb(doc)


def individual_task(doc):
    heading(doc, 'ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ')
    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=2)
    r = p.add_run('по производственной (проектно-технологической) практике')
    sf(r, bold=True, underline=True, size=14)
    p2 = doc.add_paragraph()
    pf(p2, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=12)
    r2 = p2.add_run('(вид и тип практики)')
    sf(r2, italic=True, size=12)

    para(doc, 'Для студентки [ФАМИЛИЯ ИМЯ ОТЧЕСТВО]' + '_' * 20,
         first=False, sa=0)
    p3 = doc.add_paragraph()
    pf(p3, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=6)
    r3 = p3.add_run('(Ф.И.О.)')
    sf(r3, italic=True, size=12)

    para(doc, 'Направление подготовки ', first=False, sa=0)
    p4 = doc.add_paragraph()
    pf(p4, first=False, sa=0)
    r4a = p4.add_run('09.03.02 «Информационные системы и технологии»')
    sf(r4a, underline=True, size=14)

    p5 = doc.add_paragraph()
    pf(p5, first=False, sa=0)
    r5a = p5.add_run('Направленность (профиль) подготовки ')
    sf(r5a, size=14)
    r5b = p5.add_run('«Информационные системы и сетевые технологии»')
    sf(r5b, underline=True, size=14)

    para(doc, 'Срок прохождения практики: с 04.04.2026 г. по 18.05.2026 г.',
         first=False, sa=12)

    COMPETENCIES = [
        ('ПК-1.1', 'Применяет методы и средства сборки и интеграции программных '
         'модулей и компонент, методы и средства верификации работоспособности '
         'выпусков программных продуктов, языки, утилиты и среды программирования, '
         'средства пакетного выполнения процедур'),
        ('ПК-1.2', 'Выполняет процедуры сборки программных модулей и компонент в '
         'программный продукт, производить настройки параметров программного продукта '
         'и осуществлять запуск процедур сборки'),
        ('ПК-1.3', 'Проводит оценку работоспособности программного продукта, '
         'документировать произведённые действия, выявленные проблемы и '
         'способы их устранения'),
        ('ПК-3.1', 'Решает практические задачи по созданию резервных копий БД, '
         'специальные знания по работе с установленной БД'),
        ('ПК-3.2', 'Выявляет угрозы безопасности на уровне БД, разрабатывает '
         'мероприятия по обеспечению безопасности на уровне БД'),
        ('ПК-4.1', 'Выявляет и анализирует требования к информационным системам'),
        ('ПК-4.2', 'Осуществляет сопровождение информационных систем'),
        ('ПК-5.1', 'Понимает общие принципы функционирования аппаратных, программных '
         'и программно-аппаратных средств администрируемой сети'),
        ('ПК-5.2', 'Устанавливает и настраивает программное обеспечение в '
         'соответствии с регламентами проведения профилактических работ на '
         'администрируемой инфокоммуникационной системе'),
        ('ПК-5.3', 'Управляет доступом пользователей к программно-аппаратным '
         'средствам информационных служб инфокоммуникационной системы'),
        ('ПК-6.1', 'Устанавливает и конфигурирует сетевые устройства и '
         'программное обеспечение'),
        ('ПК-6.2', 'Контролирует производительность сетевой инфраструктуры '
         'информационно-коммуникационной системы'),
        ('ПК-6.3', 'Проводит регламентные работы на сетевых устройствах и '
         'программном обеспечении информационно-коммуникационной системы'),
    ]

    for code, text in COMPETENCIES:
        p = doc.add_paragraph()
        pf(p, first=True, sa=3)
        r_code = p.add_run(f'{code} - ')
        sf(r_code, size=14)
        r_text = p.add_run(text)
        sf(r_text, size=14)

    para(doc, 'Индивидуальное задание:', first=True, sa=4, sb=6)
    para(doc,
         'Разработка адаптивного веб-сервиса мониторинга мерчандайзинга '
         'на языке Python с использованием Django 5.x и PostgreSQL. '
         'Выполнить анализ предметной области контроля выкладки товаров, '
         'определить функциональные требования, спроектировать структуру базы данных '
         'из семи взаимосвязанных моделей, реализовать регистрацию и авторизацию '
         'пользователей, создание и ведение визитов в торговые точки, '
         'фотофиксацию выкладки, чек-листы проверки SKU с AJAX-обновлением, '
         'расчёт процента выполнения плана, двухуровневую аналитику с '
         'интерактивными графиками Chart.js, провести тестирование системы '
         'и подготовить отчётные материалы.',
         sa=18)

    p6 = doc.add_paragraph()
    pf(p6, first=False, sa=0)
    r6 = p6.add_run('Руководитель практики от кафедры МТУСИ, ассистент Тришина С.В.')
    sf(r6, bold=True, size=14)
    p7 = doc.add_paragraph()
    pf(p7, first=False, sa=0)
    r7 = p7.add_run('_' * 80)
    sf(r7, size=14)
    p8 = doc.add_paragraph()
    pf(p8, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=0)
    r8 = p8.add_run('(должность, Ф.И.О., подпись)')
    sf(r8, italic=True, size=12)
    p9 = doc.add_paragraph()
    pf(p9, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=12)
    r9 = p9.add_run('04.04.2026')
    sf(r9, size=14)

    p10 = doc.add_paragraph()
    pf(p10, first=False, sa=0)
    r10a = p10.add_run('Согласовано:\nРуководитель практики от организации')
    sf(r10a, bold=True, size=14)
    r10b = p10.add_run('  ' + '_' * 30)
    sf(r10b, size=14)
    p11 = doc.add_paragraph()
    pf(p11, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=0)
    r11 = p11.add_run('(должность, Ф.И.О., подпись)')
    sf(r11, italic=True, size=12)
    p12 = doc.add_paragraph()
    pf(p12, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=12)
    r12 = p12.add_run('04.04.2026')
    sf(r12, size=14)

    p13 = doc.add_paragraph()
    pf(p13, first=False, sa=0)
    r13a = p13.add_run('Студентка [ФАМИЛИЯ ИМЯ ОТЧЕСТВО]')
    sf(r13a, bold=True, size=14)
    r13b = p13.add_run('_' * 30)
    sf(r13b, size=14)
    p14 = doc.add_paragraph()
    pf(p14, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=0)
    r14 = p14.add_run('(Ф.И.О., подпись)')
    sf(r14, italic=True, size=12)
    p15 = doc.add_paragraph()
    pf(p15, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=0)
    r15 = p15.add_run('04.04.2026')
    sf(r15, size=14)
    pb(doc)


def work_plan(doc):
    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sb=4, sa=2)
    r = p.add_run('ПЛАН (рабочий график)')
    sf(r, bold=True, size=14)

    p2 = doc.add_paragraph()
    pf(p2, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=2)
    r2 = p2.add_run('по производственной (проектно-технологической) практике')
    sf(r2, bold=True, underline=True, size=14)
    p3 = doc.add_paragraph()
    pf(p3, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=10)

    para(doc, 'Для студентки [ФАМИЛИЯ ИМЯ ОТЧЕСТВО]' + '_' * 20,
         first=False, sa=0)
    p4 = doc.add_paragraph()
    pf(p4, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=4)
    r4 = p4.add_run('( Ф.И.О.)')
    sf(r4, italic=True, size=12)

    p5 = doc.add_paragraph()
    pf(p5, first=False, sa=0)
    r5a = p5.add_run('Направление подготовки ')
    sf(r5a, size=14)
    r5b = p5.add_run('09.03.02 «Информационные системы и технологии»')
    sf(r5b, underline=True, size=14)

    p6 = doc.add_paragraph()
    pf(p6, first=False, sa=0)
    r6a = p6.add_run('Направленность (профиль) подготовки ')
    sf(r6a, size=14)
    r6b = p6.add_run('«Информационные системы и сетевые технологии»')
    sf(r6b, underline=True, size=14)

    para(doc, 'Срок прохождения практики: с 04.04.2026 г. по 18.05.2026 г.',
         first=False, sa=0)
    para(doc, 'в    Кафедра СИТиС МТУСИ' + '_' * 40, first=False, sa=0)
    p7 = doc.add_paragraph()
    pf(p7, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=8)
    r7 = p7.add_run('(наименование организации)')
    sf(r7, italic=True, size=12)

    PLAN = [
        ('04.04.2026',
         'Инструктаж, ознакомление с программой практики, анализ предметной области.',
         'ПК-1.1, ПК-4.1'),
        ('04.04.2026–17.04.2026',
         'Проектирование БД, разработка моделей данных, настройка Django-проекта и '
         'административного интерфейса.',
         'ПК-1.2, ПК-4.2'),
        ('18.04.2026–01.05.2026',
         'Реализация функционала визитов, чек-листов, загрузки фотографий, '
         'завершения визита и отчёта.',
         'ПК-1.3, ПК-3.1, ПК-5.3'),
        ('02.05.2026–12.05.2026',
         'Разработка модуля аналитики, регистрации пользователей, '
         'тестирование и исправление ошибок.',
         'ПК-3.2, ПК-5.2, ПК-6.2'),
        ('13.05.2026–18.05.2026',
         'Оформление отчёта, дневника и подготовка итоговых материалов.',
         'ПК-6.3'),
    ]

    table = doc.add_table(rows=1 + len(PLAN), cols=3)
    table.style = 'Table Grid'
    for ci, h in enumerate(['Период (дата)',
                             'Содержание практики (наименование работ)',
                             'Планируемые результаты\n(освоенные компетенции)']):
        p = table.rows[0].cells[ci].paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r = p.add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for ri, (period, content, comp) in enumerate(PLAN):
        row = table.rows[ri + 1]
        for ci, val in enumerate([period, content, comp]):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 2)
                else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            r = p.add_run(val)
            r.font.name = 'Times New Roman'; r.font.size = Pt(11)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph()

    p8 = doc.add_paragraph()
    pf(p8, first=False, sa=0)
    r8 = p8.add_run('Руководитель практики от кафедры МТУСИ, ассистент, Тришина С.В.')
    sf(r8, bold=True, size=14)
    p9 = doc.add_paragraph()
    pf(p9, first=False, sa=0)
    r9 = p9.add_run('_' * 80)
    sf(r9, size=14)
    p10 = doc.add_paragraph()
    pf(p10, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=0)
    r10 = p10.add_run('(должность, Ф.И.О., подпись)')
    sf(r10, italic=True, size=12)
    p11 = doc.add_paragraph()
    pf(p11, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=12)
    r11 = p11.add_run('18.05.2026')
    sf(r11, size=14)

    p12 = doc.add_paragraph()
    pf(p12, first=False, sa=0)
    r12a = p12.add_run('Согласовано:\nРуководитель практики от организации')
    sf(r12a, bold=True, size=14)
    r12b = p12.add_run('  ' + '_' * 30)
    sf(r12b, size=14)
    p13 = doc.add_paragraph()
    pf(p13, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=0)
    r13 = p13.add_run('(должность, Ф.И.О., подпись)')
    sf(r13, italic=True, size=12)
    p14 = doc.add_paragraph()
    pf(p14, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=12)
    r14 = p14.add_run('18.05.2026')
    sf(r14, size=14)

    p15 = doc.add_paragraph()
    pf(p15, first=False, sa=0)
    r15a = p15.add_run('Студентка [ФАМИЛИЯ ИМЯ ОТЧЕСТВО]')
    sf(r15a, bold=True, size=14)
    r15b = p15.add_run('_' * 30)
    sf(r15b, size=14)
    p16 = doc.add_paragraph()
    pf(p16, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=0)
    r16 = p16.add_run('(Ф.И.О., подпись)')
    sf(r16, italic=True, size=12)
    p17 = doc.add_paragraph()
    pf(p17, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=0)
    r17 = p17.add_run('18.05.2026')
    sf(r17, size=14)
    pb(doc)


def review_page(doc):
    heading(doc, 'Отзыв')
    p = doc.add_paragraph()
    pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=2)
    r = p.add_run('о прохождении производственной (проектно-технологической) практики')
    sf(r, bold=True, underline=True, size=14)
    p2 = doc.add_paragraph()
    pf(p2, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=10)
    r2 = p2.add_run('(вид и тип практики)')
    sf(r2, italic=True, size=12)

    para(doc, 'Студентка [ФАМИЛИЯ ИМЯ ОТЧЕСТВО]' + '_' * 20,
         first=False, sa=0)
    p3 = doc.add_paragraph()
    pf(p3, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=6)
    r3 = p3.add_run('(Ф.И.О)')
    sf(r3, italic=True, size=12)

    p4 = doc.add_paragraph()
    pf(p4, first=False, sa=0)
    r4a = p4.add_run('Направление подготовки ')
    sf(r4a, size=14)
    r4b = p4.add_run('09.03.02 «Информационные системы и технологии»')
    sf(r4b, underline=True, size=14)

    p5 = doc.add_paragraph()
    pf(p5, first=False, sa=0)
    r5a = p5.add_run('Направленность (профиль) подготовки ')
    sf(r5a, size=14)
    r5b = p5.add_run('«Информационные системы и сетевые технологии»')
    sf(r5b, underline=True, size=14)

    para(doc, 'Срок прохождения практики: с 04.04.2026 г. по 18.05.2026 г.',
         first=False, sa=0)
    para(doc, 'в    Кафедра СИТиС МТУСИ' + '_' * 40, first=False, sa=0)
    p6 = doc.add_paragraph()
    pf(p6, align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=10)
    r6 = p6.add_run('(наименование организации)')
    sf(r6, italic=True, size=12)

    para(doc,
         'Отзыв о прохождении студенткой практики (выполнение индивидуального задания, '
         'плана практики, отношение к работе, трудовая дисциплина, овладение '
         'производственными навыками, участие в научно-исследовательской, '
         'рационализаторской, общественной работе и др.).',
         italic=True, sa=4)

    para(doc,
         'В период прохождения производственной практики студентка выполнила '
         'индивидуальное задание в полном объёме. В ходе практики был проведён '
         'анализ предметной области контроля мерчандайзинга, спроектирована '
         'схема базы данных из семи моделей и разработан адаптивный веб-сервис '
         'на Python и Django 5.x. Студентка реализовала и проверила механизмы '
         'регистрации и авторизации пользователей, создания и ведения визитов, '
         'загрузки фотографий выкладки, чек-листов проверки SKU с AJAX-обновлением, '
         'расчёта процента выполнения плана, а также двухуровневой аналитики '
         'с интерактивными графиками Chart.js. За время практики студентка проявила '
         'дисциплинированность, самостоятельность, ответственность и '
         'заинтересованность в результате. План практики выполнен полностью, '
         'отчётная документация подготовлена в полном объёме, поставленные '
         'задачи решены на хорошем профессиональном уровне.',
         sa=18)

    p7 = doc.add_paragraph()
    pf(p7, first=False, sa=6)
    r7 = p7.add_run('Оценка результатов прохождения практики и выполнения индивидуального задания')
    sf(r7, bold=True, size=14)

    para(doc, 'Отлично',
         align=WD_ALIGN_PARAGRAPH.CENTER, first=False, sa=12)

    p8 = doc.add_paragraph()
    pf(p8, first=False, sa=0)
    r8a = p8.add_run('Ф.И.О., подпись ')
    sf(r8a, size=14)
    r8b = p8.add_run('_' * 30)
    sf(r8b, size=14)
    p9 = doc.add_paragraph()
    pf(p9, align=WD_ALIGN_PARAGRAPH.RIGHT, first=False, sa=0)
    r9 = p9.add_run('Печать (организации)')
    sf(r9, bold=True, size=14)
    para(doc, 'Дата    18.05.2026' + '_' * 20, first=False, sa=0)


# ─────────────────────────── main ────────────────────────────────────────────

doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

title_page(doc)
direction_page(doc)
diary_table(doc)
individual_task(doc)
work_plan(doc)
review_page(doc)

out = 'Дневник_по_практике_Мерчандайзинг.docx'
doc.save(out)
print(f'Готово: {out}')
